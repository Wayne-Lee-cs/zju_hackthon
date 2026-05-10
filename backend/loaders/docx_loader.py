import os
from docx import Document
from ..models import Textbook, Chapter


def load_docx(file_path: str, filename: str) -> Textbook:
    doc = Document(file_path)
    chapters = []
    current_chapter = None
    current_content = []

    for para in doc.paragraphs:
        if para.style.name.startswith('Heading'):
            if current_chapter:
                content = '\n'.join(current_content)
                chapters.append(Chapter(
                    chapter_id=f"ch_{len(chapters)+1}",
                    title=current_chapter,
                    page_start=0,
                    page_end=0,
                    content=content,
                    char_count=len(content)
                ))
            current_chapter = para.text
            current_content = []
        else:
            current_content.append(para.text)

    if current_chapter:
        content = '\n'.join(current_content)
        chapters.append(Chapter(
            chapter_id=f"ch_{len(chapters)+1}",
            title=current_chapter,
            page_start=0,
            page_end=0,
            content=content,
            char_count=len(content)
        ))

    if not chapters:
        content = '\n'.join([para.text for para in doc.paragraphs])
        chapters.append(Chapter(
            chapter_id="ch_1",
            title="全文",
            page_start=0,
            page_end=0,
            content=content,
            char_count=len(content)
        ))

    total_chars = sum(ch.char_count for ch in chapters)
    title = os.path.splitext(filename)[0]
    return Textbook(
        textbook_id=f"book_{hash(filename) % 10000:04d}",
        filename=filename,
        title=title,
        total_pages=0,
        total_chars=total_chars,
        chapters=chapters
    )
